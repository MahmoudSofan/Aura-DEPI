"""Format-specific document parsers (research.md §8).

Each adapter returns a :class:`ParsedDocument` (text + page_count + metadata)
or raises :class:`DocumentParseError` for unparseable input. The dispatch
function :func:`parse` selects the adapter by ``DocumentFormat``.

PDF — pdfplumber (text layer only; image-only PDFs are rejected)
DOCX — python-docx
TXT / MD — Path.read_text(utf-8) with chardet fallback
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from agents.schemas import DocumentFormat

from rag import DocumentParseError

logger = logging.getLogger("aura.rag.parsers")


@dataclass
class ParsedDocument:
    """Output of a format-specific parser."""

    text: str
    page_count: int | None = None
    parse_metadata: dict[str, Any] = field(default_factory=dict)


# ---------------------------------------------------------------------------
# Format-specific adapters.
# ---------------------------------------------------------------------------


def parse_pdf(path: Path) -> ParsedDocument:
    """Extract text from a text-based PDF using pdfplumber.

    Image-only PDFs (no extractable text on any page) raise
    :class:`DocumentParseError` per research.md §8.
    """

    try:
        import pdfplumber
    except ImportError as exc:  # pragma: no cover — dep is in pyproject
        raise DocumentParseError(f"pdfplumber not available: {exc}") from exc

    pages_text: list[str] = []
    try:
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                pages_text.append(page_text)
    except Exception as exc:  # noqa: BLE001 — narrow via DocumentParseError
        raise DocumentParseError(f"pdfplumber failed to open {path.name}: {exc}") from exc

    joined = "\n\n".join(pt for pt in pages_text if pt.strip())
    if not joined.strip():
        raise DocumentParseError(f"{path.name}: no extractable text (image-only or empty PDF)")

    return ParsedDocument(
        text=joined,
        page_count=len(pages_text),
        parse_metadata={"parser": "pdfplumber", "pages": len(pages_text)},
    )


def parse_docx(path: Path) -> ParsedDocument:
    """Extract text from a DOCX file using python-docx."""

    try:
        from docx import Document as DocxDocument
    except ImportError as exc:  # pragma: no cover
        raise DocumentParseError(f"python-docx not available: {exc}") from exc

    try:
        doc = DocxDocument(str(path))
    except Exception as exc:  # noqa: BLE001
        raise DocumentParseError(f"{path.name}: docx open failed: {exc}") from exc

    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    text = "\n\n".join(paragraphs)

    if not text.strip():
        raise DocumentParseError(f"{path.name}: docx produced no text")

    return ParsedDocument(
        text=text,
        page_count=None,
        parse_metadata={"parser": "python-docx", "paragraph_count": len(paragraphs)},
    )


def _read_text_with_fallback(path: Path) -> str:
    """Decode bytes as UTF-8 first, fall back to chardet detection."""

    raw = path.read_bytes()
    if not raw:
        raise DocumentParseError(f"{path.name}: file is empty")

    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        pass

    try:
        import chardet
    except ImportError:  # pragma: no cover
        return raw.decode("utf-8", errors="replace")

    detected = chardet.detect(raw)
    encoding = detected.get("encoding") or "utf-8"
    try:
        return raw.decode(encoding, errors="replace")
    except (UnicodeDecodeError, LookupError) as exc:
        raise DocumentParseError(f"{path.name}: undecodable text: {exc}") from exc


def parse_text(path: Path) -> ParsedDocument:
    """Read a plain-text file with UTF-8 + chardet fallback."""

    text = _read_text_with_fallback(path)
    if not text.strip():
        raise DocumentParseError(f"{path.name}: text file has no non-whitespace content")
    return ParsedDocument(
        text=text,
        page_count=None,
        parse_metadata={"parser": "text"},
    )


def parse_markdown(path: Path) -> ParsedDocument:
    """Read a markdown file. Behaves like :func:`parse_text` for v1."""

    text = _read_text_with_fallback(path)
    if not text.strip():
        raise DocumentParseError(f"{path.name}: markdown file has no non-whitespace content")
    return ParsedDocument(
        text=text,
        page_count=None,
        parse_metadata={"parser": "markdown"},
    )


# ---------------------------------------------------------------------------
# Dispatch.
# ---------------------------------------------------------------------------


_DISPATCH = {
    "pdf": parse_pdf,
    "docx": parse_docx,
    "txt": parse_text,
    "md": parse_markdown,
}


def parse(path: Path, format: DocumentFormat) -> ParsedDocument:
    """Dispatch to the right parser based on ``format``."""

    parser = _DISPATCH.get(format)
    if parser is None:
        raise DocumentParseError(f"unsupported format: {format!r}")
    return parser(path)


__all__ = [
    "ParsedDocument",
    "parse",
    "parse_docx",
    "parse_markdown",
    "parse_pdf",
    "parse_text",
]
